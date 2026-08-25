"""
Report generation - DOCX (python-docx) and PDF (reportlab), both driven by
the same "template settings" dict so a company's branding is defined once
and reused across formats.

The report is structured into exactly 7 sections, written so a
non-technical reader (a manager, HR, legal counsel) can follow the case
without needing to understand raw logs:
  1. Executive Summary
  2. Scope & Evidence Reviewed
  3. Timeline of Key Events        (every doubtful/notable timestamp, no cap)
  4. Detailed Findings              (with POC images/files embedded inline)
  5. Indicator of Compromise Summary
  6. Risk Assessment
  7. Recommendations
followed by a non-numbered Appendix A (chain-of-custody evidence manifest).

template settings dict:
{
  "company_name": str, "header_text": str, "footer_text": str,
  "primary_color": "#RRGGBB",              # accent color for tables/rules
  "heading_style": "Same" | "Distinct",
  "body_style": "Formal" | "Plain",
  "include_bullets": bool, "include_tables": bool,
  "heading_font": str, "heading_color": "#RRGGBB", "heading_size": int,
  "subheading_font": str, "subheading_color": "#RRGGBB", "subheading_size": int,
  "body_font": str, "body_color": "#RRGGBB", "body_size": int,
}
heading_font/subheading_font/body_font are keys into FONT_CHOICES below -
each maps to a real DOCX font name and the nearest built-in ReportLab PDF
font (ReportLab ships only Helvetica/Times/Courier without embedding a
TTF, so PDF font choice is "family" rather than exact typeface).
"""

import os
import json
from datetime import datetime

from analyzers.narrative import build_story, burst_sentence, aggregate_events, reconstruct_finding_timestamps
from analyzers.m365_classifier import SCHEMA_LABELS as _M365_SCHEMA_LABELS
from analyzers.vehicle_classifier import SCHEMA_LABELS as _VEHICLE_SCHEMA_LABELS
from core import db

SCHEMA_LABELS = {**_M365_SCHEMA_LABELS, **_VEHICLE_SCHEMA_LABELS}

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors as rl_colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None


FONT_CHOICES = {
    "Calibri (Sans-serif)": dict(docx="Calibri", pdf="Helvetica", pdf_bold="Helvetica-Bold"),
    "Arial (Sans-serif)": dict(docx="Arial", pdf="Helvetica", pdf_bold="Helvetica-Bold"),
    "Verdana (Sans-serif)": dict(docx="Verdana", pdf="Helvetica", pdf_bold="Helvetica-Bold"),
    "Times New Roman (Serif)": dict(docx="Times New Roman", pdf="Times-Roman", pdf_bold="Times-Bold"),
    "Georgia (Serif)": dict(docx="Georgia", pdf="Times-Roman", pdf_bold="Times-Bold"),
    "Courier New (Monospace)": dict(docx="Courier New", pdf="Courier", pdf_bold="Courier-Bold"),
}
DEFAULT_FONT_CHOICE = "Calibri (Sans-serif)"

IMAGE_EXT = (".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff")


def _hex_to_rgb(hex_color):
    hex_color = (hex_color or "#1F2937").lstrip("#")
    if len(hex_color) != 6:
        hex_color = "1F2937"
    return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))


def _font(settings, key):
    return FONT_CHOICES.get(settings.get(key, DEFAULT_FONT_CHOICE), FONT_CHOICES[DEFAULT_FONT_CHOICE])


STATUS_ORDER = ["Verified", "Needs Further Investigation", "Unverified", "False Positive"]


def _sorted_findings(findings):
    return sorted(
        findings,
        key=lambda f: (STATUS_ORDER.index(f["verification_status"]) if f["verification_status"] in STATUS_ORDER else 99,
                        -f["confidence"]),
    )


def _confidence_label(f):
    """Email findings don't get a meaningful automated confidence score (point 7) - shown as '-' instead of a number."""
    return "-" if f["type"] == "email" else f"{f['confidence']:.2f}"


def _reconstructed_timestamps(findings, events):
    """
    Timestamp reconstruction: for every finding, work out when it was
    first/last actually observed by matching it back to the Investigation
    Story timeline (the same events/bursts stage 3 builds), rather than
    just listing the indicator with no sense of when it happened.
    Returns {finding_id: "First Seen str" / "Last Seen str" / approx flag}.
    """
    bursts, _undated = aggregate_events(events) if events else ([], 0)
    out = {}
    for f in findings:
        first_seen, last_seen, approx = reconstruct_finding_timestamps(f["value"], bursts)
        out[f["id"]] = dict(first_seen=first_seen, last_seen=last_seen, approximate=approx)
    return out


def _fmt_ts(dt, approx):
    if not dt:
        return ""
    label = dt.strftime("%Y-%m-%d %H:%M:%S")
    return f"~{label}" if approx else label


def _time_label(item):
    label = item["start"].strftime("%Y-%m-%d %H:%M:%S")
    if item["end"] != item["start"]:
        label += f" - {item['end'].strftime('%H:%M:%S')}"
    if item.get("approximate"):
        label += " (approx.)"
    return label


def _object_value(objects, prefix):
    prefix_l = prefix.lower() + ":"
    for o in objects or []:
        if str(o).lower().startswith(prefix_l):
            return str(o).split(":", 1)[1]
    return None


# --------------------------------------------------------------------------
# Section-content builders shared by DOCX and PDF
# --------------------------------------------------------------------------

def _evidence_scope_rows(evidence, events):
    """
    Section 2 (Scope & Evidence Reviewed): for every evidence file, work out
    a human label for what kind of export it is (auto-detected M365/Entra
    schema when recognized, otherwise the analyst's chosen category) and
    which mailbox(es) it's actually about - so a reader sees e.g.
    "M365 Unified Audit Log (XLSX)" / "jdoe@corp-internal.local" instead of just
    a filename.
    """
    events_by_file = {}
    for e in events or []:
        events_by_file.setdefault(e.get("source_file"), []).append(e)

    rows = []
    for ev in evidence:
        fname = ev["filename"]
        file_events = events_by_file.get(fname, [])
        ext = os.path.splitext(fname)[1].lstrip(".").upper()

        schema_counts = {}
        for e in file_events:
            s = e.get("schema")
            if s:
                schema_counts[s] = schema_counts.get(s, 0) + 1
        if schema_counts:
            top_schema = max(schema_counts, key=schema_counts.get)
            label = SCHEMA_LABELS.get(top_schema, top_schema)
        else:
            label = ev.get("category") or "Evidence file"
        type_label = f"{label} ({ext})" if ext else label

        mailboxes = set()
        for e in file_events:
            actor = e.get("actor")
            if actor and "@" in str(actor):
                mailboxes.add(str(actor))
            mbox = _object_value(e.get("objects"), "Mailbox")
            if mbox and "@" in str(mbox):
                mailboxes.add(str(mbox))
            vin = _object_value(e.get("objects"), "VIN")
            if vin:
                mailboxes.add(str(vin))
        mailboxes = sorted(mailboxes)
        mailbox_label = ", ".join(mailboxes[:3]) + (", ..." if len(mailboxes) > 3 else "")

        rows.append(dict(filename=fname, type_label=type_label, mailbox=mailbox_label or "-"))
    return rows


def _ioc_ip_rows(notable_bursts):
    """Section 5 IP table: IP / Activity (what happened) / Where (location, if known) - built only from bursts tied to a detected pattern, not routine noise."""
    by_ip = {}
    for b in notable_bursts:
        ip = b.get("source_ip")
        if not ip:
            continue
        entry = by_ip.setdefault(ip, {"activities": set(), "locations": set()})
        entry["activities"].add(burst_sentence(b))
        loc = _object_value(b.get("objects"), "Location")
        if loc:
            entry["locations"].add(loc)
    rows = []
    for ip in sorted(by_ip):
        data = by_ip[ip]
        rows.append(dict(
            ip=ip,
            activity="; ".join(sorted(data["activities"]))[:400],
            where=", ".join(sorted(data["locations"])) or "Unknown",
        ))
    return rows


def _ioc_inbox_rule_rows(notable_bursts):
    """Section 5 Inbox Rules table."""
    rows = []
    for b in notable_bursts:
        if b.get("category") != "Mailbox Rule":
            continue
        rows.append(dict(
            mailbox=b.get("actor") or "",
            rule_name=_object_value(b.get("objects"), "RuleName") or "",
            forward_to=_object_value(b.get("objects"), "ForwardTo") or "",
            outcome=b.get("outcome") or "",
            when=_time_label(b),
        ))
    return rows


def _ioc_other_findings(findings):
    """Section 5 catch-all: domains/hashes/URLs worth calling out as indicators, beyond IPs and inbox rules already tabulated above."""
    types = ("domain", "hash_sha256", "hash_sha1", "hash_md5", "url", "crypto_btc", "crypto_eth")
    return [f for f in findings if f["type"] in types and f["verification_status"] != "False Positive"]


_RECOMMENDATION_RULES = [
    ("brute-force", "immediate", "Force a password reset and enable multi-factor authentication (MFA) for the affected account(s) immediately."),
    ("brute-force", "followup", "Review sign-in activity from the same source IP(s) against other accounts to rule out a wider password-spray attempt."),
    ("beacon", "immediate", "Isolate the affected host from the network and block the destination IP/domain at the firewall or proxy."),
    ("beacon", "followup", "Run a full antivirus/EDR scan on the affected host and check it for persistence mechanisms (scheduled tasks, startup entries, new services)."),
    ("malware", "immediate", "Isolate the affected host, preserve a forensic image, and run a full antivirus/EDR scan before returning it to service."),
    ("lookalike domain", "immediate", "Block the lookalike domain at the email gateway/web proxy and alert staff who may have received or clicked it."),
    ("lookalike domain", "followup", "Monitor for newly registered domains that mimic your company's domain to catch repeat attempts early."),
    ("outbound mail", "immediate", "Quarantine or recall the outbound message(s) if not yet delivered, and notify the recipient(s) not to act on it."),
    ("outbound", "followup", "Review the sending account's recent mailbox activity for other signs of compromise (new rules, unfamiliar sign-ins)."),
    ("inbox rule", "immediate", "Remove or disable the suspicious inbox rule and reset the mailbox owner's password."),
    ("inbox rule", "followup", "Audit all inbox rules across the tenant for the same forward/auto-delete pattern - this is a common post-compromise persistence tactic."),
    ("oauth", "immediate", "Revoke the OAuth application's consent/access tokens and remove it from the tenant if it isn't a recognized, approved app."),
    ("consent", "immediate", "Revoke the OAuth application's consent/access tokens and remove it from the tenant if it isn't a recognized, approved app."),
    ("oauth", "followup", "Require admin approval for future third-party application consent grants to prevent unreviewed app access."),
    ("delegation", "immediate", "Remove the unauthorized mailbox delegation/permission grant."),
    ("delegation", "followup", "Audit all mailbox delegations tenant-wide for other unauthorized or stale grants."),
    ("travel", "immediate", "Contact the account owner to confirm whether both sign-in locations are legitimate; reset credentials immediately if not."),
    ("travel", "followup", "Enable or tighten Conditional Access location-based and impossible-travel policies."),
    ("off-hours", "followup", "Confirm with the account owner whether the off-hours activity was expected; adjust access policies if it wasn't."),
]


def build_recommendations(insights):
    """
    Section 7: turns whichever pattern types were actually detected in
    THIS case into concrete "do this now" / "do this as follow-up" actions,
    instead of generic boilerplate that doesn't reflect what happened.
    """
    immediate, followup, seen = [], [], set()
    for ins in insights:
        title_l = ins["title"].lower()
        matched = False
        for keyword, bucket, text in _RECOMMENDATION_RULES:
            if keyword in title_l and text not in seen:
                seen.add(text)
                (immediate if bucket == "immediate" else followup).append(text)
                matched = True
        if not matched and ins.get("severity") == "High":
            generic = f"Investigate and remediate the root cause of: {ins['title']}."
            if generic not in seen:
                seen.add(generic)
                immediate.append(generic)

    baseline = [
        "Preserve all evidence, logs, and this report under chain-of-custody pending closure of the investigation.",
        "Review this report with legal/compliance to determine any breach-notification obligations.",
    ]
    for b in baseline:
        if b not in seen:
            followup.append(b)
    return immediate, followup


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------

def generate_docx(case, findings, evidence, settings, output_path, include_unverified=True, events=None, insights=None):
    events = events or []
    insights = insights or []
    r, g, b = _hex_to_rgb(settings.get("primary_color"))
    accent_rgb = RGBColor(r, g, b)
    heading_font = _font(settings, "heading_font")
    subheading_font = _font(settings, "subheading_font")
    body_font = _font(settings, "body_font")
    heading_rgb = RGBColor(*_hex_to_rgb(settings.get("heading_color") or settings.get("primary_color")))
    subheading_rgb = RGBColor(*_hex_to_rgb(settings.get("subheading_color") or settings.get("primary_color")))
    body_rgb = RGBColor(*_hex_to_rgb(settings.get("body_color") or "#000000"))
    heading_size = int(settings.get("heading_size") or 20)
    subheading_size = int(settings.get("subheading_size") or 14)
    body_size = int(settings.get("body_size") or 11)

    doc = Document()

    # Apply body font/size/color as the document default so every plain
    # paragraph and table cell picks it up without per-call styling.
    normal = doc.styles["Normal"]
    normal.font.name = body_font["docx"]
    normal.font.size = Pt(body_size)
    normal.font.color.rgb = body_rgb

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = settings.get("header_text") or settings.get("company_name") or ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    footer.text = settings.get("footer_text") or f"Confidential - {case['case_id']}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def heading(text, level=1):
        h = doc.add_heading(text, level=level)
        font_choice = heading_font if level == 1 else subheading_font
        color = heading_rgb if level == 1 else subheading_rgb
        size = heading_size if level == 1 else subheading_size
        if settings.get("heading_style") == "Same":
            size = body_size + 2
        for run in h.runs:
            run.font.name = font_choice["docx"]
            run.font.size = Pt(size)
            run.font.color.rgb = color
        return h

    def para(text, bullet=False):
        return doc.add_paragraph(text, style="List Bullet" if (bullet and settings.get("include_bullets", True)) else None)

    def table_with_header(cols, style="Light Grid Accent 1"):
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = style
        hdr = table.rows[0].cells
        for i, col in enumerate(cols):
            hdr[i].text = col
        return table

    title = doc.add_heading("Vehicle & Cyber Forensic Investigation Report", level=0)
    for run in title.runs:
        run.font.color.rgb = heading_rgb
        run.font.name = heading_font["docx"]
    sub = doc.add_paragraph()
    sub.add_run(f"Case: {case['name']}  |  Case ID: {case['case_id']}").bold = True
    doc.add_paragraph(f"Prepared by: {settings.get('company_name') or case.get('company') or ''}")
    doc.add_paragraph(f"Reporter: {case.get('reporter') or ''}")
    doc.add_paragraph(f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    for e in events:
        e.setdefault("_idx", e.get("id"))
    story = build_story(events, insights, case) if events else None

    verified = [f for f in findings if f["verification_status"] == "Verified"]
    to_show = findings if include_unverified else verified

    # --- 1. Executive Summary ---
    heading("1. Executive Summary", 1)
    para(case.get("story") or "No case overview provided.")
    para(
        f"Classification: {case.get('case_type') or 'N/A'}   |   "
        f"Severity: {case.get('severity') or 'N/A'}   |   "
        f"Confidentiality: {case.get('confidentiality') or 'N/A'}"
    )
    if story:
        para(story["executive_summary"])
        high_count = len([i for i in insights if i["severity"] == "High"])
        if high_count:
            para(f"This report flags {high_count} high-severity issue(s) requiring immediate attention - see Section 6 (Risk Assessment) and Section 7 (Recommendations).")
    else:
        para("No timeline evidence was processed for this case yet.")

    # --- 2. Scope & Evidence Reviewed ---
    heading("2. Scope & Evidence Reviewed", 1)
    para(f"{len(evidence)} evidence file(s) were reviewed for this investigation:")
    scope_rows = _evidence_scope_rows(evidence, events)
    if settings.get("include_tables", True):
        t = table_with_header(["Filename", "Type", "Associated Mailbox / VIN"])
        for row in scope_rows:
            cells = t.add_row().cells
            cells[0].text, cells[1].text, cells[2].text = row["filename"], row["type_label"], row["mailbox"]
    else:
        for row in scope_rows:
            para(f"{row['filename']} - {row['type_label']} - mailbox: {row['mailbox']}", bullet=True)

    # --- 3. Timeline of Key Events ---
    heading("3. Timeline of Key Events", 1)
    para(
        "Every doubtful or notable timestamp found across the evidence is listed below, oldest first. "
        "Each entry shows exactly what the log recorded, followed in plain language by what it means and why it matters "
        "- no raw log reading required."
    )
    if story and story["timeline_story"]:
        for point in story["timeline_story"]:
            p = doc.add_paragraph()
            p.add_run(f"{_time_label(point)} - ").bold = True
            p.add_run(point["technical_line"])
            for ins in point["explanations"]:
                wp = doc.add_paragraph(style="List Bullet 2" if settings.get("include_bullets", True) else None)
                wp.add_run(f"[{ins['severity']}] What this means: ").bold = True
                wp.add_run(ins["description"])
                if ins.get("impact"):
                    ip_p = doc.add_paragraph(style="List Bullet 2" if settings.get("include_bullets", True) else None)
                    ip_p.add_run("Impact: ").bold = True
                    ip_p.add_run(ins["impact"])
    else:
        para("No notable/doubtful timestamps were detected in the evidence reviewed.")

    # --- 4. Detailed Findings ---
    heading("4. Detailed Findings", 1)
    ts_map = _reconstructed_timestamps(to_show, events)
    para(
        "First Seen / Last Seen are reconstructed from the timeline above by matching each indicator back to the log "
        "events that reference it. A \"~\" prefix means the timestamp is approximate; blank means no timestamped event "
        "could be matched. Email-type findings don't get an automated confidence score (\"-\") since deduplication "
        "confidence there isn't a meaningful risk signal."
    )
    if settings.get("include_tables", True):
        t = table_with_header(["Type", "Value", "First Seen", "Last Seen", "Confidence", "Status", "Notes"])
        for f in _sorted_findings(to_show):
            ts = ts_map.get(f["id"], {})
            row = t.add_row().cells
            row[0].text = f["type"]
            row[1].text = f["value"]
            row[2].text = _fmt_ts(ts.get("first_seen"), ts.get("approximate"))
            row[3].text = _fmt_ts(ts.get("last_seen"), ts.get("approximate"))
            row[4].text = _confidence_label(f)
            row[5].text = f["verification_status"]
            row[6].text = f.get("verification_notes") or ""
    else:
        for f in _sorted_findings(to_show):
            ts = ts_map.get(f["id"], {})
            when = _fmt_ts(ts.get("first_seen"), ts.get("approximate"))
            p = para(
                f"[{f['type']}] {f['value']} - confidence {_confidence_label(f)}, status: {f['verification_status']}"
                + (f", first seen: {when}" if when else ""),
                bullet=True,
            )
            if f.get("verification_notes"):
                para(f"    Notes: {f['verification_notes']}")

    with_poc = [f for f in to_show if db.get_poc_paths(f["id"])]
    if with_poc:
        heading("Proof of Concept (POC) attachments", 2)
        for f in with_poc:
            para(f"[{f['type']}] {f['value']}").runs[0].bold = True
            for poc_path in db.get_poc_paths(f["id"]):
                if poc_path.lower().endswith(IMAGE_EXT):
                    try:
                        doc.add_picture(poc_path, width=Inches(4.5))
                    except Exception:
                        para(f"(could not embed image: {os.path.basename(poc_path)})")
                else:
                    para(f"📄 Attached file: {os.path.basename(poc_path)}", bullet=True)

    # --- 5. Indicator of Compromise Summary ---
    heading("5. Indicator of Compromise Summary", 1)
    notable_bursts = story["notable_bursts"] if story else []
    ip_rows = _ioc_ip_rows(notable_bursts)
    if ip_rows:
        heading("IP Addresses", 2)
        if settings.get("include_tables", True):
            t = table_with_header(["IP Address", "Activity (what happened)", "Where"])
            for row in ip_rows:
                cells = t.add_row().cells
                cells[0].text, cells[1].text, cells[2].text = row["ip"], row["activity"], row["where"]
        else:
            for row in ip_rows:
                para(f"{row['ip']} - {row['activity']} - {row['where']}", bullet=True)
    else:
        para("No notable IP addresses were tied to a detected pattern.")

    rule_rows = _ioc_inbox_rule_rows(notable_bursts)
    if rule_rows:
        heading("Inbox Rules", 2)
        if settings.get("include_tables", True):
            t = table_with_header(["Mailbox", "Rule Name", "Forward To", "Outcome", "When"])
            for row in rule_rows:
                cells = t.add_row().cells
                cells[0].text, cells[1].text, cells[2].text, cells[3].text, cells[4].text = (
                    row["mailbox"], row["rule_name"], row["forward_to"], row["outcome"], row["when"])
        else:
            for row in rule_rows:
                para(f"{row['when']} - {row['mailbox']} - rule '{row['rule_name']}' forwarding to {row['forward_to']}", bullet=True)

    other = _ioc_other_findings(to_show)
    if other:
        heading("Other Indicators", 2)
        for f in other:
            para(f"[{f['type']}] {f['value']} - status: {f['verification_status']}", bullet=True)

    # --- 6. Risk Assessment ---
    heading("6. Risk Assessment", 1)
    if insights:
        if settings.get("include_tables", True):
            t = table_with_header(["Finding", "Status", "Impact"])
            for ins in insights:
                cells = t.add_row().cells
                cells[0].text = ins["title"]
                cells[1].text = ins["severity"]
                cells[2].text = ins.get("impact") or ins.get("description") or ""
        else:
            for ins in insights:
                para(f"[{ins['severity']}] {ins['title']}: {ins.get('impact') or ins.get('description') or ''}", bullet=True)
    else:
        para("No automated patterns were detected - a manual risk assessment is recommended.")

    # --- 7. Recommendations ---
    heading("7. Recommendations", 1)
    immediate, followup = build_recommendations(insights)
    heading("Immediate Actions", 2)
    if immediate:
        for r_ in immediate:
            para(r_, bullet=True)
    else:
        para("No high-severity findings required immediate action at the time of writing.")
    heading("Follow-up / Hardening Recommendations", 2)
    for r_ in followup:
        para(r_, bullet=True)

    # --- Appendix A: Chain of custody ---
    heading("Appendix A: Chain of Custody - Evidence Manifest", 1)
    t = table_with_header(["Filename", "SHA-256", "Category", "Uploaded At"], style="Light List Accent 1")
    for e in evidence:
        cells = t.add_row().cells
        cells[0].text = e["filename"]
        cells[1].text = e["sha256"]
        cells[2].text = e.get("category") or ""
        cells[3].text = e["uploaded_at"]

    doc.save(output_path)
    return output_path


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------

def generate_pdf(case, findings, evidence, settings, output_path, include_unverified=True, events=None, insights=None):
    events = events or []
    insights = insights or []
    r, g, b = _hex_to_rgb(settings.get("primary_color"))
    primary_color = rl_colors.Color(r / 255, g / 255, b / 255)
    heading_font = _font(settings, "heading_font")
    subheading_font = _font(settings, "subheading_font")
    body_font = _font(settings, "body_font")
    heading_color = rl_colors.Color(*[c / 255 for c in _hex_to_rgb(settings.get("heading_color") or settings.get("primary_color"))])
    subheading_color = rl_colors.Color(*[c / 255 for c in _hex_to_rgb(settings.get("subheading_color") or settings.get("primary_color"))])
    body_color = rl_colors.Color(*[c / 255 for c in _hex_to_rgb(settings.get("body_color") or "#000000")])
    heading_size = int(settings.get("heading_size") or 20)
    subheading_size = int(settings.get("subheading_size") or 14)
    body_size_pt = int(settings.get("body_size") or 11)

    styles = getSampleStyleSheet()
    h1_style = ParagraphStyle("H1Custom", parent=styles["Heading1"], textColor=heading_color,
                                fontName=heading_font["pdf_bold"], fontSize=heading_size, leading=heading_size + 4)
    h2_style = ParagraphStyle("H2Custom", parent=styles["Heading2"], textColor=subheading_color,
                                fontName=subheading_font["pdf_bold"], fontSize=subheading_size, leading=subheading_size + 3)
    title_style = ParagraphStyle("TitleCustom", parent=styles["Title"], textColor=heading_color,
                                   fontName=heading_font["pdf_bold"])
    body_style = ParagraphStyle("BodyCustom", parent=styles["BodyText"], textColor=body_color,
                                  fontName=body_font["pdf"], fontSize=body_size_pt, leading=body_size_pt + 4)
    bold_body_style = ParagraphStyle("BodyBoldCustom", parent=body_style, fontName=body_font["pdf_bold"])

    header_text = settings.get("header_text") or settings.get("company_name") or ""
    footer_text = settings.get("footer_text") or f"Confidential - {case['case_id']}"

    def on_page(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.drawString(0.75 * inch, letter[1] - 0.5 * inch, header_text)
        canvas.drawCentredString(letter[0] / 2, 0.5 * inch, f"{footer_text}  |  Page {doc_.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(output_path, pagesize=letter, topMargin=0.9 * inch, bottomMargin=0.9 * inch)
    flow = []

    def bullet_table(cols, rows, col_widths):
        data = [cols] + rows
        t = Table(data, repeatRows=1, colWidths=col_widths)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), primary_color),
            ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
            ("FONTNAME", (0, 0), (-1, -1), body_font["pdf"]),
            ("FONTSIZE", (0, 0), (-1, -1), max(7, body_size_pt - 3)),
            ("GRID", (0, 0), (-1, -1), 0.5, rl_colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return t

    flow.append(Paragraph("Vehicle & Cyber Forensic Investigation Report", title_style))
    flow.append(Spacer(1, 12))
    flow.append(Paragraph(f"<b>Case:</b> {case['name']}  |  <b>Case ID:</b> {case['case_id']}", body_style))
    flow.append(Paragraph(f"Prepared by: {settings.get('company_name') or case.get('company') or ''}", body_style))
    flow.append(Paragraph(f"Reporter: {case.get('reporter') or ''}", body_style))
    flow.append(Paragraph(f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body_style))
    flow.append(Spacer(1, 16))

    for e in events:
        e.setdefault("_idx", e.get("id"))
    story = build_story(events, insights, case) if events else None

    verified = [f for f in findings if f["verification_status"] == "Verified"]
    to_show = findings if include_unverified else verified

    # --- 1. Executive Summary ---
    flow.append(Paragraph("1. Executive Summary", h1_style))
    flow.append(Paragraph(case.get("story") or "No case overview provided.", body_style))
    flow.append(Paragraph(
        f"Classification: {case.get('case_type') or 'N/A'} | Severity: {case.get('severity') or 'N/A'} | "
        f"Confidentiality: {case.get('confidentiality') or 'N/A'}", body_style))
    if story:
        flow.append(Paragraph(story["executive_summary"], body_style))
        high_count = len([i for i in insights if i["severity"] == "High"])
        if high_count:
            flow.append(Paragraph(
                f"This report flags {high_count} high-severity issue(s) requiring immediate attention - see "
                "Section 6 (Risk Assessment) and Section 7 (Recommendations).", body_style))
    else:
        flow.append(Paragraph("No timeline evidence was processed for this case yet.", body_style))
    flow.append(Spacer(1, 10))

    # --- 2. Scope & Evidence Reviewed ---
    flow.append(Paragraph("2. Scope & Evidence Reviewed", h1_style))
    flow.append(Paragraph(f"{len(evidence)} evidence file(s) were reviewed for this investigation:", body_style))
    scope_rows = _evidence_scope_rows(evidence, events)
    if settings.get("include_tables", True) and scope_rows:
        rows = [[row["filename"], row["type_label"], row["mailbox"]] for row in scope_rows]
        flow.append(bullet_table(["Filename", "Type", "Associated Mailbox / VIN"], rows,
                                   [2 * inch, 2.4 * inch, 2.1 * inch]))
    else:
        for row in scope_rows:
            flow.append(Paragraph(f"&bull; {row['filename']} - {row['type_label']} - mailbox: {row['mailbox']}", body_style))
    flow.append(Spacer(1, 10))

    # --- 3. Timeline of Key Events ---
    flow.append(Paragraph("3. Timeline of Key Events", h1_style))
    flow.append(Paragraph(
        "Every doubtful or notable timestamp found across the evidence is listed below, oldest first, each followed "
        "in plain language by what it means and why it matters.", body_style))
    if story and story["timeline_story"]:
        for point in story["timeline_story"]:
            flow.append(Paragraph(f"<b>{_time_label(point)}</b> - {point['technical_line']}", bold_body_style))
            for ins in point["explanations"]:
                flow.append(Paragraph(f"&bull; <b>[{ins['severity']}] What this means:</b> {ins['description']}", body_style))
                if ins.get("impact"):
                    flow.append(Paragraph(f"&bull; <b>Impact:</b> {ins['impact']}", body_style))
            flow.append(Spacer(1, 4))
    else:
        flow.append(Paragraph("No notable/doubtful timestamps were detected in the evidence reviewed.", body_style))
    flow.append(Spacer(1, 8))

    # --- 4. Detailed Findings ---
    flow.append(Paragraph("4. Detailed Findings", h1_style))
    ts_map = _reconstructed_timestamps(to_show, events)
    flow.append(Paragraph(
        "First Seen / Last Seen are reconstructed from the timeline above. A \"~\" prefix means approximate; blank "
        "means no timestamped event could be matched. Email-type findings show \"-\" for confidence.", body_style))
    if settings.get("include_tables", True):
        data_rows = []
        for f in _sorted_findings(to_show):
            ts = ts_map.get(f["id"], {})
            data_rows.append([f["type"], f["value"], _fmt_ts(ts.get("first_seen"), ts.get("approximate")),
                               _fmt_ts(ts.get("last_seen"), ts.get("approximate")),
                               _confidence_label(f), f["verification_status"], f.get("verification_notes") or ""])
        flow.append(bullet_table(["Type", "Value", "First Seen", "Last Seen", "Confidence", "Status", "Notes"],
                                   data_rows,
                                   [0.65 * inch, 1.5 * inch, 0.9 * inch, 0.9 * inch, 0.55 * inch, 1.0 * inch, 1.0 * inch]))
    else:
        for f in _sorted_findings(to_show):
            ts = ts_map.get(f["id"], {})
            when = _fmt_ts(ts.get("first_seen"), ts.get("approximate"))
            flow.append(Paragraph(
                f"&bull; <b>[{f['type']}]</b> {f['value']} - confidence {_confidence_label(f)}, status: {f['verification_status']}"
                + (f", first seen: {when}" if when else ""), body_style))
    flow.append(Spacer(1, 8))

    with_poc = [f for f in to_show if db.get_poc_paths(f["id"])]
    if with_poc:
        flow.append(Paragraph("Proof of Concept (POC) attachments", h2_style))
        for f in with_poc:
            flow.append(Paragraph(f"<b>[{f['type']}] {f['value']}</b>", body_style))
            for poc_path in db.get_poc_paths(f["id"]):
                if poc_path.lower().endswith(IMAGE_EXT) and PILImage is not None:
                    try:
                        with PILImage.open(poc_path) as im:
                            w, h = im.size
                        max_w = 4.5 * inch
                        scale = min(1.0, max_w / w) if w else 1.0
                        flow.append(RLImage(poc_path, width=w * scale, height=h * scale))
                    except Exception:
                        flow.append(Paragraph(f"(could not embed image: {os.path.basename(poc_path)})", body_style))
                else:
                    flow.append(Paragraph(f"&bull; Attached file: {os.path.basename(poc_path)}", body_style))
            flow.append(Spacer(1, 6))

    # --- 5. Indicator of Compromise Summary ---
    flow.append(Paragraph("5. Indicator of Compromise Summary", h1_style))
    notable_bursts = story["notable_bursts"] if story else []
    ip_rows = _ioc_ip_rows(notable_bursts)
    if ip_rows:
        flow.append(Paragraph("IP Addresses", h2_style))
        if settings.get("include_tables", True):
            rows = [[row["ip"], row["activity"], row["where"]] for row in ip_rows]
            flow.append(bullet_table(["IP Address", "Activity (what happened)", "Where"], rows,
                                       [1.2 * inch, 3.6 * inch, 1.7 * inch]))
        else:
            for row in ip_rows:
                flow.append(Paragraph(f"&bull; {row['ip']} - {row['activity']} - {row['where']}", body_style))
    else:
        flow.append(Paragraph("No notable IP addresses were tied to a detected pattern.", body_style))
    flow.append(Spacer(1, 6))

    rule_rows = _ioc_inbox_rule_rows(notable_bursts)
    if rule_rows:
        flow.append(Paragraph("Inbox Rules", h2_style))
        if settings.get("include_tables", True):
            rows = [[row["mailbox"], row["rule_name"], row["forward_to"], row["outcome"], row["when"]] for row in rule_rows]
            flow.append(bullet_table(["Mailbox", "Rule Name", "Forward To", "Outcome", "When"], rows,
                                       [1.3 * inch, 1.0 * inch, 1.6 * inch, 0.7 * inch, 1.4 * inch]))
        else:
            for row in rule_rows:
                flow.append(Paragraph(f"&bull; {row['when']} - {row['mailbox']} - rule '{row['rule_name']}' forwarding to {row['forward_to']}", body_style))
        flow.append(Spacer(1, 6))

    other = _ioc_other_findings(to_show)
    if other:
        flow.append(Paragraph("Other Indicators", h2_style))
        for f in other:
            flow.append(Paragraph(f"&bull; [{f['type']}] {f['value']} - status: {f['verification_status']}", body_style))
    flow.append(Spacer(1, 8))

    # --- 6. Risk Assessment ---
    flow.append(Paragraph("6. Risk Assessment", h1_style))
    if insights:
        if settings.get("include_tables", True):
            rows = [[ins["title"], ins["severity"], (ins.get("impact") or ins.get("description") or "")] for ins in insights]
            flow.append(bullet_table(["Finding", "Status", "Impact"], rows, [1.8 * inch, 0.8 * inch, 3.9 * inch]))
        else:
            for ins in insights:
                flow.append(Paragraph(f"&bull; <b>[{ins['severity']}] {ins['title']}:</b> {ins.get('impact') or ins.get('description') or ''}", body_style))
    else:
        flow.append(Paragraph("No automated patterns were detected - a manual risk assessment is recommended.", body_style))
    flow.append(Spacer(1, 8))

    # --- 7. Recommendations ---
    flow.append(Paragraph("7. Recommendations", h1_style))
    immediate, followup = build_recommendations(insights)
    flow.append(Paragraph("Immediate Actions", h2_style))
    if immediate:
        for r_ in immediate:
            flow.append(Paragraph(f"&bull; {r_}", body_style))
    else:
        flow.append(Paragraph("No high-severity findings required immediate action at the time of writing.", body_style))
    flow.append(Paragraph("Follow-up / Hardening Recommendations", h2_style))
    for r_ in followup:
        flow.append(Paragraph(f"&bull; {r_}", body_style))

    # --- Appendix A: Chain of custody ---
    flow.append(PageBreak())
    flow.append(Paragraph("Appendix A: Chain of Custody - Evidence Manifest", h1_style))
    rows = [[e["filename"], e["sha256"], e.get("category") or "", e["uploaded_at"]] for e in evidence]
    flow.append(bullet_table(["Filename", "SHA-256", "Category", "Uploaded At"], rows,
                               [1.6 * inch, 2.6 * inch, 1 * inch, 1.5 * inch]))

    doc.build(flow, onFirstPage=on_page, onLaterPages=on_page)
    return output_path
