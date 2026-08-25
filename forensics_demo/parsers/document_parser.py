"""Parsers for PDF and DOCX evidence."""

import pdfplumber
import docx
from core.models import ParsedChunk


def parse_pdf(file_path: str, rel_path: str, evidence_type: str = "document"):
    chunks = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            for j, line in enumerate(text.splitlines(), start=1):
                if line.strip():
                    chunks.append(ParsedChunk(
                        text=line,
                        source_file=rel_path,
                        location=f"page {i}, line {j}",
                        evidence_type=evidence_type,
                    ))
            if not text.strip():
                # No extractable text layer -> likely a scanned page.
                chunks.append(ParsedChunk(
                    text="",
                    source_file=rel_path,
                    location=f"page {i} (no text layer - consider OCR)",
                    evidence_type=evidence_type,
                ))
    return chunks


def parse_docx(file_path: str, rel_path: str, evidence_type: str = "document"):
    chunks = []
    d = docx.Document(file_path)
    for i, para in enumerate(d.paragraphs, start=1):
        if para.text.strip():
            chunks.append(ParsedChunk(
                text=para.text,
                source_file=rel_path,
                location=f"paragraph {i}",
                evidence_type=evidence_type,
            ))
    for ti, table in enumerate(d.tables, start=1):
        for ri, row in enumerate(table.rows, start=1):
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                chunks.append(ParsedChunk(
                    text=row_text,
                    source_file=rel_path,
                    location=f"table {ti}, row {ri}",
                    evidence_type=evidence_type,
                ))
    return chunks
